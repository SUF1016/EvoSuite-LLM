from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "26项目报告.docx"
BACKUP = ROOT / "26项目报告_补充前备份.docx"
TEMP_OUTPUT = ROOT / "26项目报告_补充完善.docx"


def set_run_font(run, font_name: str = "宋体", size: int | None = 11, bold: bool | None = None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_text(doc: Document, text: str, bold: bool = False, align=None) -> Paragraph:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return paragraph


def add_section_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True)


def add_code(doc: Document, title: str, code: str) -> None:
    add_text(doc, title, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(code.strip())
    set_run_font(run, font_name="Consolas", size=8)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def delete_after(doc: Document, anchor: Paragraph) -> None:
    body = doc._body._element
    children = list(body)
    deleting = False
    sect_pr = None
    for child in children:
        if child.tag.endswith("sectPr"):
            sect_pr = child
            continue
        if deleting:
            body.remove(child)
        if child is anchor._p:
            deleting = True
    if sect_pr is not None:
        try:
            body.remove(sect_pr)
        except ValueError:
            pass
        body.append(sect_pr)


PIPELINE_CODE = """
baseline = evaluate_project(self.project_dir, self.config, run_dir, "evosuite-baseline")
business_rules = {
    target: self._extract_business_rules(target, run_dir)
    for target in targets
}

for iteration in range(1, int(self.config.get("iterations", 1)) + 1):
    for target in targets:
        messages = build_enhancement_messages(
            self.project_dir,
            target,
            current_pit,
            current_coverage,
            self.config.get("prompt", {}),
            business_rules=business_rules.get(target, ""),
        )
        response = self.llm.complete(messages)
        files = extract_java_files(response, fallback_relative_path=fallback)
        generated_paths = write_generated_files(self.project_dir, files)
        test_result = run_maven_goal(self.project_dir, self.config["maven"], "test_goal")
        if not test_result.ok:
            generated_paths = self._repair(target, generated_paths, test_result.output, run_dir)

    evaluation = evaluate_project(self.project_dir, self.config, run_dir, f"llm-iteration-{iteration}")
"""


MUTATION_CODE = """
def classify_mutation(mutation: Mutation) -> str:
    text = f"{mutation.mutator} {mutation.description}".lower()
    if "conditionalsboundary" in text or "changed conditional boundary" in text:
        return "boundary-condition"
    if "negateconditionals" in text or "negated conditional" in text:
        return "branch-polarity"
    if "mathmutator" in text or "increments" in text:
        return "numeric-formula"
    if "return" in text:
        return "return-oracle"
    return "general"
"""


PROMPT_CODE = """
Business rules to turn into semantic oracles:
{rule_context}

Survived or uncovered PIT mutants:
{mutants}

{mutant_guidance}

Task:
1. Add or improve tests that kill the highest-value survived mutants.
2. Assert exact CheckoutResult fields, rejection reasons, and BigDecimal values.
3. Boundary mutants need threshold tests; return mutants need direct assertions.
4. Keep tests deterministic, readable, and under 6 focused test methods.
"""


ASSERTION_CODE = """
@property
def assertion_strength_score(self) -> float:
    if self.test_methods == 0:
        return 0.0
    semantic_bonus = min(self.semantic_assertions * 1.5, 15)
    weak_penalty = min(self.weak_assertions * 1.0, 15)
    return round(
        min(100.0, self.assertion_strength_per_test * 12 + semantic_bonus - weak_penalty),
        2,
    )
"""


def add_metric_table(doc: Document) -> None:
    headers = ["指标", "EvoSuite baseline", "LLM iteration 1", "LLM iteration 2", "分析"]
    rows = [
        ["Test OK", "yes", "no", "yes", "测试必须通过才可作为有效结果；iteration 1 失败，iteration 2 可用。"],
        ["Line Coverage", "88.99%", "88.99%", "92.20%", "越高越好；iteration 2 有提升。"],
        ["Branch Coverage", "85.34%", "85.34%", "85.34%", "越高越好；三组持平，说明收益不主要来自新增分支覆盖。"],
        ["Mutation Score", "69.44%", "69.44%", "94.44%", "越高越好；这是本项目最关键指标，iteration 2 显著提升。"],
        ["Assertions", "29", "74", "64", "不是简单越高越好；必须结合测试是否通过、是否稳定、断言是否有语义。"],
        ["Assertion Strength", "39.1", "76.3", "68.0", "一般越高越好，但仍是启发式指标，不能单独替代 mutation score。"],
        ["Flaky Failures", "0", "2", "0", "越低越好；iteration 2 与 baseline 一样稳定。"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=9)


def add_body(doc: Document) -> None:
    add_text(
        doc,
        "本项目实现的是一个面向 Java 单元测试自动生成与增强的完整实验框架，项目名称为“基于变异测试反馈的大语言模型增强 EvoSuite 单元测试生成方法”。项目的基本思想是：不直接修改 EvoSuite 内核，而是在 EvoSuite 外层构建一个 Python 自动化编排器，让 EvoSuite 负责生成初始 JUnit 测试，让大语言模型负责根据源码语义、业务规则和变异测试反馈补充更强的测试 oracle，最后再通过 Maven、JUnit、JaCoCo 和 PIT 进行自动验证和量化评价。"
    )
    add_text(
        doc,
        "项目主体代码位于 mglet 目录。该目录中的 Python 文件负责完成实验流程控制、配置读取、外部命令调用、EvoSuite 运行、JaCoCo 覆盖率解析、PIT 变异测试解析、prompt 构造、LLM 调用、生成测试文件写入、失败测试修复以及实验报告输出。被测试对象位于 java-target 目录，是一个 Maven Java 订单结算项目。该项目包含 CheckoutCalculator、CouponPolicy、Coupon、CustomerProfile、OrderItem 等类，覆盖会员等级、优惠券、折扣封顶、免运费、税费、过期校验、首单限制和异常输入等规则。选择该实验对象的原因是：这些业务逻辑具有明显的边界条件和语义 oracle，能够更好地体现 LLM 相比纯搜索式测试生成的优势。"
    )
    add_text(
        doc,
        "完整运行时，框架首先调用 tools/evosuite-1.2.0.jar 为 order.CheckoutCalculator 和 order.CouponPolicy 生成 EvoSuite seed tests，并导出到 java-target/src/test/java/order。随后框架执行 Maven 测试，确认 seed tests 可以编译并稳定运行。之后 JaCoCo 统计 line coverage 和 branch coverage，PIT 生成 mutations.xml 并标记 killed、survived、no coverage 等变异体状态。框架会把 survived mutants 进一步分类，例如边界条件变异、分支取反变异、数值公式变异、返回值变异等，并把分类结果和业务规则一起放入 prompt。LLM 生成增强测试后，框架再次执行 mvn test；如果测试失败，则把错误日志交给 LLM 进入 repair loop。最终，框架再次运行 JaCoCo 与 PIT，输出 summary.md 和 results.json。"
    )
    add_text(
        doc,
        "本项目还实现了 Assertion Strength Score。传统测试统计通常只看断言数量，但断言数量并不总是代表测试质量，因为重复断言、脆弱断言和无意义断言也会增加数量。本项目的断言强度指标会更关注是否验证 BigDecimal 金额、结果对象字段、异常信息、优惠券拒绝原因等语义信息。因此，它可以作为 mutation score 之外的辅助指标，用于说明 LLM 是否生成了更有业务含义的 oracle。"
    )

    add_section_title(doc, "本项目所参考/基于的论文或项目来源")
    add_text(
        doc,
        "本项目主要基于 EvoSuite 官方开源项目。EvoSuite 是面向 Java 的自动单元测试生成工具，其核心思想是使用搜索式算法生成 JUnit 测试，并以覆盖率等目标作为优化方向。EvoSuite 能够快速探索输入空间，自动构造对象、调用方法并生成回归式断言，因此适合作为自动测试生成的 seed 工具。"
    )
    add_text(
        doc,
        "项目评估部分基于 JaCoCo 和 PIT/Pitest。JaCoCo 用于统计 line coverage 和 branch coverage，PIT 用于进行 mutation testing。与覆盖率相比，mutation score 更能体现测试是否真正具备缺陷检测能力，因为它会检查测试能否杀死人为注入的小变异。"
    )
    add_text(
        doc,
        "LLM 工作流借鉴了 ChatUniTest 一类生成-验证-修复思想，但本项目并不是让 LLM 从零生成测试，而是采用 EvoSuite seed + LLM enhancement 的混合方式。这样可以降低 LLM 幻觉和依赖缺失风险，同时保留 LLM 对业务语义、边界条件和强断言的补充能力。"
    )

    add_section_title(doc, "本项目在参考论文或项目的基础上做了哪些改动？")
    add_text(
        doc,
        "第一，项目将 EvoSuite 从独立工具转化为自动化流程中的 seed generator。EvoSuite 生成初始测试后，框架不会直接把结果作为最终答案，而是继续进行编译验证、覆盖率分析、变异测试分析和 LLM 增强。"
    )
    add_text(
        doc,
        "第二，项目引入变异算子类型感知的 prompt。PIT 输出 survived mutants 后，框架会根据 mutator 和 description 对变异体分类。例如 changed conditional boundary 属于 boundary-condition，negated conditional 属于 branch-polarity，return value mutation 属于 return-oracle。分类后的结果会指导 LLM 选择不同测试策略。"
    )
    add_text(
        doc,
        "第三，项目引入业务规则增强的测试 oracle 生成。框架会先从目标类和同包源码中提取业务规则，例如 VIP 折扣封顶、优惠券过期、最低消费门槛、非叠加优惠券替换规则、免运费条件和地区税率等，再要求 LLM 将这些规则转化为具体断言。"
    )
    add_text(
        doc,
        "第四，项目实现 Assertion Strength Score，用于弥补“只统计断言数量”的不足。该指标会给精确金额断言、结果字段断言、异常原因断言更高权重，对 assertNotNull 这类弱断言给较低权重，因此能更好地描述测试 oracle 的强弱。"
    )
    add_text(
        doc,
        "第五，项目加入自动验证和修复循环。LLM 输出的测试只有通过 Maven 编译、JUnit 运行和多次稳定性检查后，才会进入最终评价。如果生成测试失败，系统会把错误日志反馈给 LLM 修复。真实实验中的 iteration 1 就显示：即使断言数量和断言强度较高，只要测试未通过，就不能作为有效结果。"
    )

    add_section_title(doc, "核心代码片段")
    add_code(doc, "代码片段 1：pipeline 中的核心实验闭环。", PIPELINE_CODE)
    add_text(
        doc,
        "这段代码体现了项目的主流程：先评估 EvoSuite baseline，再抽取业务规则，随后按迭代轮次调用 LLM 生成增强测试。如果测试运行失败，则进入 repair；每轮结束后再次调用 evaluator 计算覆盖率、变异分数和静态指标。"
    )
    add_code(doc, "代码片段 2：PIT 变异算子类型分类。", MUTATION_CODE)
    add_text(
        doc,
        "这段代码对应本项目的第一个创新点。它把 PIT 原始 mutator 转换为更适合 prompt 使用的语义类别，使 LLM 不只是看到一串 survived mutants，而是能知道应该补边界测试、补分支极性测试还是补返回值断言。"
    )
    add_code(doc, "代码片段 3：业务规则与 mutation feedback 共同进入 prompt。", PROMPT_CODE)
    add_text(
        doc,
        "这段 prompt 约束体现了项目的第二个创新点。LLM 不从零自由发挥，而是在业务规则、survived mutants 和变异算子策略共同约束下生成少量、定向、可验证的测试。"
    )
    add_code(doc, "代码片段 4：Assertion Strength Score 的核心计算。", ASSERTION_CODE)
    add_text(
        doc,
        "这段代码体现了项目的第三个创新点。断言强度分数由每个测试方法平均断言强度、语义断言奖励和弱断言惩罚共同决定，最终限制在 0 到 100 之间。"
    )

    add_section_title(doc, "实验结果")
    add_text(
        doc,
        "本次真实 API 实验使用 qwen3.6-plus，运行目录为 runs/course-project/20260521-213542，目标类为 order.CheckoutCalculator 和 order.CouponPolicy，迭代轮数为 2。实验先使用 EvoSuite seed tests 作为 baseline，再执行两轮 mutation-guided LLM enhancement。"
    )
    add_metric_table(doc)

    add_section_title(doc, "指标含义说明")
    add_text(
        doc,
        "Test OK 表示测试套件是否能够通过 Maven 编译和 JUnit 运行。该指标是最基础的有效性门槛。如果 Test OK 为 no，则即使断言数量很多、断言强度分数很高，也不能作为最终有效结果。"
    )
    add_text(
        doc,
        "Line Coverage 表示被执行到的代码行比例，通常越高越好。它可以说明测试是否执行了更多代码，但不能直接说明断言是否有效。Branch Coverage 表示分支条件的覆盖比例，也通常越高越好，但同样不能单独代表 oracle 强度。"
    )
    add_text(
        doc,
        "Mutation Score 是本项目最重要的指标。PIT 会对生产代码注入小变异，例如修改边界条件、取反条件、替换返回值等。如果测试能够发现变异导致的行为变化，该变异体就被 killed；如果测试仍然通过，该变异体就是 survived。Mutation score 越高，说明测试越能发现潜在缺陷。"
    )
    add_text(
        doc,
        "Assertions 表示断言数量。它不是绝对越高越好，只能作为辅助指标。断言多可能说明 oracle 更丰富，也可能只是重复断言、脆弱断言或无意义断言。Iteration 1 的断言数为 74，高于 iteration 2 的 64，但 iteration 1 测试未通过且存在 flaky failures，因此不能说明 iteration 1 更好。"
    )
    add_text(
        doc,
        "Assertion Strength 是启发式断言强度指标，一般越高越好，但仍不能脱离 Test OK 和 Mutation Score 单独使用。它更偏向奖励精确金额断言、结果字段断言、异常原因断言和业务语义断言。Iteration 2 的 Assertion Strength 为 68.0，高于 baseline 的 39.1，说明最终可用测试的 oracle 强度相较 EvoSuite 原始测试明显增强。"
    )
    add_text(
        doc,
        "Flaky Failures 表示重复运行中出现失败的次数，越低越好。Iteration 1 有 2 次 flaky failures，说明它虽然生成了更多断言，但稳定性不足。Iteration 2 和 baseline 都为 0，说明最终增强测试是稳定可用的。"
    )

    add_section_title(doc, "实验结果分析")
    add_text(
        doc,
        "从整体结果看，iteration 2 并不是每个数值都严格高于 iteration 1，但它是唯一同时满足测试通过、无 flaky、mutation score 显著提升、断言强度高于 baseline 的最终有效结果。因此，实验结论不能简单写成“iteration 2 全面优于所有指标”，而应写成“iteration 2 是最终可用测试套件中效果最好的方案”。"
    )
    add_text(
        doc,
        "与 EvoSuite baseline 相比，iteration 2 的 line coverage 从 88.99% 提升到 92.20%，mutation score 从 69.44% 提升到 94.44%，Assertion Strength 从 39.1 提升到 68.0，并保持 Test OK 为 yes 和 Flaky Failures 为 0。这说明 LLM 增强没有破坏测试可运行性，同时显著增强了测试发现缺陷的能力。"
    )
    add_text(
        doc,
        "Branch coverage 在三组中均为 85.34%，说明本项目的主要收益不是来自分支覆盖率扩大。对于 EvoSuite 这类 coverage-driven 工具，覆盖率本来已经较高，继续提高 branch coverage 的空间有限。LLM 的贡献主要体现在业务 oracle 和边界断言上，也就是让已有覆盖路径上的测试更容易杀死变异体。"
    )
    add_text(
        doc,
        "Iteration 1 是一个反例：它生成了更多断言，Assertions 达到 74，Assertion Strength 达到 76.3，但 Test OK 为 no，且 Flaky Failures 为 2。因此，单独追求断言数量或断言强度并不可靠。自动化验证和 repair loop 是必要的，因为 LLM 可能生成看似合理但不稳定或无法通过的测试。"
    )
    add_text(
        doc,
        "Iteration 2 经过反馈修复后，虽然断言数量和断言强度低于 iteration 1，但测试能够稳定通过，并把 survived mutants 从 baseline 的 20 个减少到 4 个，no coverage mutants 从 2 个减少到 0 个，killed mutants 从 50 个提升到 68 个。这说明最终保留下来的断言更有效、更稳定，也更符合软件工程中“可运行、可复现、可验证”的要求。"
    )

    add_section_title(doc, "项目总结（不少于300字）")
    add_section_title(doc, "简述自己在项目中的收获（比如难点的解决、运用的软件工程工具和技术、软件开发与AI的结合等）")
    add_text(
        doc,
        "本项目完成了一个较完整的软件工程实验原型。它不仅包含代码生成，还包含工具链集成、自动验证、失败修复、指标评估和报告输出。通过该项目可以看到，自动化测试生成并不是单一模型或单一工具能够完全解决的问题。EvoSuite 能够高效生成 seed tests，适合探索输入空间和提升覆盖率；LLM 能够根据源码和业务规则补充更有语义的测试 oracle；PIT 能够指出哪些变异体没有被杀死，为下一轮生成提供目标；Maven/JUnit 则负责保证测试真正可编译、可运行。把这些工具串起来，才形成了一个有工程意义的闭环。"
    )
    add_text(
        doc,
        "项目开发过程中最大的收获之一是认识到指标之间存在优先级。覆盖率越高通常越好，但覆盖率高不代表测试强；断言数量越多也不一定越好，因为失败测试、重复断言和脆弱断言都可能增加数量。Mutation score 更能体现测试缺陷检测能力，但也必须建立在测试可稳定运行的基础上。因此，最终评价不能只看单一数值，而要综合 Test OK、Flaky Failures、Coverage、Mutation Score 和 Assertion Strength。"
    )
    add_text(
        doc,
        "项目中的另一个收获是对 AI 编程工具风险的认识更加具体。LLM 在生成测试时确实能根据业务语义补充边界用例和强断言，但它也可能猜测不存在的枚举常量、生成过长导致代码块截断、忽略测试框架约束，甚至产生不稳定测试。因此，本项目没有直接信任 LLM 输出，而是通过 Maven 编译、JUnit 运行、PIT 评估和 repair loop 对其进行约束。这种做法更符合高级软件工程课程中强调的质量保证思想。"
    )
    add_text(
        doc,
        "在工程实现方面，本项目还解决了多个实际问题。例如 EvoSuite 1.2.0 在 JDK 17 下容易受到 Java module 反射限制，因此最终使用 JDK 8；中文路径下 JaCoCo agent 输出可能失败，因此脚本将当前项目临时映射到 M: 盘；EvoSuite 生成的部分时间相关测试可能 flaky，因此需要固定日期；Qwen 模型在长 prompt 下容易超时，因此关闭 thinking 并增加 timeout 与 retries。这些问题虽然不是算法本身，但正是工程项目中非常重要的可复现性和稳定性问题。"
    )
    add_text(
        doc,
        "总体来看，本项目证明了 EvoSuite seed tests 与 LLM enhancement 的结合具有实际价值。真实实验中，iteration 2 在保持测试稳定通过的前提下，把 mutation score 从 69.44% 提升到 94.44%，说明增强后的测试更能发现代码缺陷。这一结果支持项目的核心假设：LLM 的价值不是替代 EvoSuite，而是在 EvoSuite 已经生成高覆盖率测试的基础上，利用业务语义理解能力增强 oracle 和边界断言。"
    )

    add_section_title(doc, "简述未来工作展望。")
    add_text(
        doc,
        "未来工作可以从三个方向继续扩展。第一，扩大实验对象，从当前订单结算项目扩展到 Defects4J、GitBug-Java 或 Apache Commons 系列项目，以验证方法在真实开源项目上的泛化能力。第二，增加更严格的对比组，例如纯 LLM 从零生成测试、ChatUniTest 默认方法、只使用业务规则但不使用 mutation feedback 的消融实验等。第三，进一步改进 LLM 输出控制和测试筛选机制，例如对同一个 survived mutant 生成多个候选测试，保留能稳定杀死变异体且断言最强的一组，同时自动删除重复或脆弱测试。"
    )


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(REPORT, BACKUP)
    doc = Document(REPORT)
    anchor = find_paragraph(doc, "项目功能介绍（不少于500字）")
    delete_after(doc, anchor)
    add_body(doc)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name not in {"Consolas"}:
                set_run_font(run)

    doc.save(TEMP_OUTPUT)
    shutil.copy2(TEMP_OUTPUT, REPORT)
    print(f"updated={REPORT}")
    print(f"backup={BACKUP}")
    print(f"draft={TEMP_OUTPUT}")


if __name__ == "__main__":
    main()
