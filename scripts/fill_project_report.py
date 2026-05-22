from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "26项目报告.docx"
OUTPUT = ROOT / "26项目报告_已填写.docx"


def set_run_font(run, font_name: str = "宋体", size=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = size
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run)


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def insert_lines_after(paragraph: Paragraph, lines: list[str]) -> None:
    cursor = paragraph
    for line in lines:
        cursor = insert_after(cursor, line)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


PROJECT_INTRO = [
    "本项目实现了一个“基于变异测试反馈的大语言模型增强 EvoSuite 单元测试生成方法”原型系统。项目的核心目标不是重新实现 EvoSuite，也不是直接让大语言模型从零生成测试，而是把两类方法的优势组合起来：先利用 EvoSuite 的搜索式测试生成能力为 Java 类自动生成可运行的 JUnit seed tests，再利用大语言模型根据源码语义、业务规则、JaCoCo 覆盖率和 PIT survived mutants 对测试进行二次增强。EvoSuite 擅长探索输入空间和快速获得较高覆盖率，但生成的断言通常偏向回归式 oracle，测试命名和可读性也较弱；大语言模型具备一定语义理解能力，但直接生成测试时容易出现 API 猜测、编译错误和不稳定测试。因此，本项目采用“生成-验证-修复-再评估”的工程流程，把 LLM 生成结果放入 Maven 编译、JUnit 运行和 PIT 变异测试中进行自动校验。",
    "当前实验对象是一个订单结算 Java Maven 项目，路径为 java-target。该项目包含 CheckoutCalculator、CouponPolicy、Coupon、CustomerProfile、OrderItem 等类，覆盖会员折扣、优惠券有效性、优惠叠加、折扣封顶、免运费、地区税率、过期校验、首单限制和异常输入等业务规则。相比简单工具类，这类场景更能体现 LLM 的价值，因为很多测试 oracle 不能只靠覆盖率判断，而需要理解金额边界、业务条件和返回字段之间的关系。例如 VIP 折扣为 10% 且封顶 40.00，非叠加优惠券只有在优于会员折扣时才替换折扣，过期优惠券必须返回 EXPIRED，数字商品应免运费，EU/US 地区税率不同。这些规则被抽取后会进入 prompt，引导 LLM 生成更强的语义断言。",
    "系统主体由 Python 包 mglet 实现。EvoSuite Runner 负责调用 tools/evosuite-1.2.0.jar 生成 JUnit 测试；Evaluator 负责调用 Maven、JaCoCo 和 PIT 并解析 XML 报告；Prompt Builder 负责构造包含源码、已有测试、业务规则和变异体信息的提示词；LLM Enhancer 通过阿里云 DashScope 的 OpenAI-compatible 接口调用 qwen3.6-plus；Validation Loop 会在生成测试后执行 mvn test，失败时把编译或运行日志反馈给 LLM 进入 repair；Reporter 最终输出 results.json 和 summary.md。项目还提供 Assertion Strength Score，用于衡量断言是否真正验证业务结果，而不是只统计断言数量。"
]

SOURCES = [
    "本项目主要基于 EvoSuite 官方开源项目。EvoSuite 是面向 Java 的自动单元测试生成工具，使用搜索式/遗传算法生成 JUnit 测试，并以覆盖率为主要优化目标。本项目没有修改 EvoSuite 内核，而是把 EvoSuite 作为命令行 seed test 生成器使用。",
    "项目评估基于 PIT/Pitest 和 JaCoCo。JaCoCo 用于统计 line coverage 和 branch coverage，PIT 用于进行 mutation testing，并输出 killed、survived、no coverage 等变异体信息。本项目把 mutation score 作为主要效果指标，因为单纯覆盖率不能充分反映测试 oracle 强度。",
    "LLM 工作流借鉴了 ChatUniTest 一类“生成-验证-修复”思想，但本项目的生成入口不同：不是让 LLM 从零生成测试，而是先由 EvoSuite 提供稳定 seed tests，再由 LLM 根据 survived mutants 和业务规则进行增强。"
]

CHANGES = [
    "1. 引入 EvoSuite-seeded LLM 增强流程：原始 EvoSuite 只负责生成 seed tests，本项目在其后增加 LLM enhancement 阶段，让模型基于已有测试补充更具语义的断言和边界用例。",
    "2. 引入变异算子类型感知 prompt：系统会解析 PIT survived mutants，并按 boundary-condition、branch-polarity、numeric-formula、return-oracle 等类型分类。不同类型的变异体会映射到不同测试策略，例如边界条件变异要求补充阈值上下界用例，返回值变异要求直接断言布尔值、字符串或结果字段。",
    "3. 引入业务规则增强的测试 oracle 生成：系统会先从目标类及同包源码中抽取业务规则，再要求 LLM 将规则转化为 JUnit 断言。订单金额、折扣上限、过期原因、税率、免运费条件等规则会被显式写入 prompt。",
    "4. 引入 Assertion Strength Score：除断言数量外，项目会根据断言类型、是否验证 BigDecimal 金额、是否验证结果字段、是否验证异常/原因字符串等因素计算断言强度分数，用来补充 coverage 和 mutation score。",
    "5. 增加自动验证和修复机制：LLM 输出的测试必须经过 Maven 编译和 JUnit 运行。若失败，系统会把错误日志反馈给 LLM 修复，避免把不可编译或不稳定的测试计入最终结果。",
    "6. 工程上支持 JDK 8、Maven、本地 EvoSuite jar、DashScope OpenAI-compatible API、中文路径下的 M: 临时映射运行，以及 JSON/Markdown 实验报告输出。"
]

RESULTS = [
    "本次真实 API 实验使用 qwen3.6-plus，运行目录为 runs/course-project/20260521-213542，目标类为 order.CheckoutCalculator 和 order.CouponPolicy，迭代轮数为 2。实验先使用 EvoSuite seed tests 作为 baseline，再执行两轮 mutation-guided LLM enhancement。",
    "实验指标如下：",
    "方法\t测试是否通过\tLine Coverage\tBranch Coverage\tMutation Score\t断言数\tAssertion Strength\tFlaky Failures",
    "EvoSuite baseline\t通过\t88.99%\t85.34%\t69.44%\t29\t39.1\t0",
    "LLM iteration 1\t未通过\t88.99%\t85.34%\t69.44%\t74\t76.3\t2",
    "LLM iteration 2\t通过\t92.20%\t85.34%\t94.44%\t64\t68.0\t0",
    "从结果看，第一轮 LLM 生成了更多语义断言，但仍存在测试不稳定或失败问题，因此不能直接计入最终有效结果。这体现了 validation loop 的必要性。第二轮经过反馈修正后，测试全部通过，flaky failures 降为 0，mutation score 从 69.44% 提升到 94.44%，killed mutants 从 50 个提升到 68 个，survived mutants 从 20 个下降到 4 个，no coverage mutants 从 2 个下降到 0 个。Line coverage 从 88.99% 提升到 92.20%，branch coverage 保持 85.34%。这说明本项目的主要改进不只是覆盖更多代码，而是显著增强了测试 oracle 对变异缺陷的检测能力。",
    "从断言质量看，Assertion Strength Score 从 39.1 提升到 68.0，可读性分数从 37.6 提升到 49.1。最终生成的 LLM 增强测试文件包括 java-target/src/test/java/order/CheckoutCalculatorLLMEnhancedTest.java 和 java-target/src/test/java/order/CouponPolicyLLMEnhancedTest.java。"
]

SUMMARY = [
    "本项目完成了一个可运行、可评估、可复现实验的课程项目原型。通过把 EvoSuite、LLM、Maven、JaCoCo 和 PIT 串联起来，项目形成了较完整的软件工程自动化流程：自动生成测试、自动运行测试、自动提取覆盖率和变异测试反馈、自动构造 prompt、自动调用模型增强测试、自动修复失败测试，并最终输出对比报告。相比只展示 LLM 生成代码，本项目更强调工程闭环和可量化结果。真实 API 实验表明，直接依赖第一轮 LLM 输出并不可靠，模型可能生成失败或不稳定测试；但在变异测试反馈和验证修复循环下，第二轮结果显著提升了 mutation score，并保持测试稳定通过。这与项目假设一致：LLM 的价值不在于替代 EvoSuite，而在于利用语义理解能力补足 EvoSuite 在业务 oracle 和强断言方面的不足。",
    "项目过程中遇到的主要难点包括 Java/JDK 版本兼容、中文路径下 JaCoCo agent 输出失败、EvoSuite 生成测试中的时间相关 flaky 问题、LLM 输出过长导致代码块截断、模型猜测不存在枚举常量、以及 API 非流式请求超时。针对这些问题，项目分别采用 JDK 8 运行 EvoSuite、M: 盘路径映射、固定日期稳定 seed test、限制 prompt 输出规模、把完整同包源码放入 prompt、关闭 Qwen thinking 并增加请求超时和重试等方式解决。通过这些处理，我对软件工程中的自动化验证、实验复现、工具链集成和 AI 代码生成风险有了更直接的认识。"
]

FUTURE = [
    "未来工作可以从三个方向继续扩展。第一，扩大实验对象，从当前小型订单结算项目扩展到 Defects4J、GitBug-Java 或 Apache Commons 系列项目，验证方法在真实开源项目上的泛化能力。第二，引入更细粒度的 flaky test filtering 和 test smell detection，对 EvoSuite 与 LLM 生成测试进行自动清洗。第三，改进 prompt 和搜索策略，例如把 survived mutants 按优先级分批处理、对同一变异体生成多个候选测试并择优保留、加入纯 LLM baseline 和 ChatUniTest baseline，以形成更严格的消融实验。"
]


def main() -> None:
    doc = Document(TEMPLATE)

    replacements = {
        "学    院：": "学    院：          （请填写）",
        "学    系：": "学    系：          （请填写）",
        "专    业：": "专    业：          （请填写）",
        "学生姓名（学号）：": "学生姓名（学号）：  （请填写）",
        "年  月  日": "2026 年 5 月 21 日",
    }
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if stripped in replacements:
            set_paragraph_text(paragraph, replacements[stripped])

    insert_lines_after(
        find_paragraph(doc, "1.代码   2.数据集  3.用法Readme文件"),
        [
            "项目 GitHub 链接：待上传 GitHub 后填写。",
            "当前本地项目路径：D:\\研\\研一\\研一下课程\\高级软件工程\\项目。",
            "代码包括 mglet Python 框架、java-target Java 实验对象、configs 配置、scripts 运行脚本和 tests 单元测试；实验数据和结果保存在 runs/course-project/20260521-213542；用法说明见 README.md 和 PROJECT_STRUCTURE.md。",
        ],
    )
    insert_lines_after(find_paragraph(doc, "项目功能介绍（不少于500字）"), PROJECT_INTRO)
    insert_lines_after(find_paragraph(doc, "本项目所参考/基于的论文或项目来源"), SOURCES)
    insert_lines_after(find_paragraph(doc, "本项目在参考论文或项目的基础上做了哪些改动？"), CHANGES)
    insert_lines_after(find_paragraph(doc, "实验结果"), RESULTS)
    insert_lines_after(
        find_paragraph(doc, "简述自己在项目中的收获（比如难点的解决、运用的软件工程工具和技术、软件开发与AI的结合等）"),
        SUMMARY,
    )
    insert_lines_after(find_paragraph(doc, "简述未来工作展望。"), FUTURE)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
